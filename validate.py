#!/usr/bin/env python3
"""
THI Solar Dashboard - Test & Validation Script

Single command to run all tests and generate a comprehensive report.

Usage:
    python validate.py                    # Generate all formats (HTML, Markdown, DOCX)
    python validate.py --html-only        # Generate HTML only
    python validate.py --quick            # Skip slow checks (Streamlit smoke test)
"""

from __future__ import annotations

import argparse
import html
import io
import os
import platform
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Sequence

from docx import Document


@dataclass
class CheckResult:
    name: str
    status: str  # PASS/FAIL/WARN/SKIP
    details: str


STATUS_ORDER = {"FAIL": 0, "WARN": 1, "SKIP": 2, "PASS": 3}


def _status_badge(status: str) -> str:
    status = (status or "").upper()
    if status == "PASS":
        return "✅ PASS"
    if status == "FAIL":
        return "❌ FAIL"
    if status == "WARN":
        return "⚠️ WARN"
    if status == "SKIP":
        return "⏭️ SKIP"
    return status


def _overall_status(results: Sequence[CheckResult]) -> str:
    statuses = [r.status.upper() for r in results]
    if any(s == "FAIL" for s in statuses):
        return "FAIL"
    if any(s == "WARN" for s in statuses):
        return "WARN"
    if statuses and all(s == "SKIP" for s in statuses):
        return "SKIP"
    return "PASS"


def _run(cmd: list[str], *, cwd: Path, timeout_s: int = 60) -> tuple[int, str]:
    """Run a command and return (exit_code, combined_output)."""
    proc = subprocess.run(
        cmd,
        cwd=str(cwd),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        timeout=timeout_s,
    )
    return proc.returncode, proc.stdout


def _pkg_version(module_name: str) -> str:
    try:
        from importlib.metadata import version
        return version(module_name)
    except Exception:
        return "(not installed)"


def _shorten(text: str, *, limit: int = 6000) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n... (truncated) ...\n"


def _first_line(text: str, *, limit: int = 140) -> str:
    line = (text or "").strip().splitlines()[0] if (text or "").strip() else ""
    if len(line) <= limit:
        return line
    return line[:limit].rstrip() + "…"


def _try_run_git(project_root: Path, args: list[str]) -> tuple[bool, str]:
    if not (project_root / ".git").exists():
        return False, "(not a git repo)"
    code, out = _run(["git", *args], cwd=project_root, timeout_s=20)
    if code != 0:
        return False, (out.strip() or "(git command failed)")
    return True, out.strip()


def check_environment() -> CheckResult:
    versions = {
        "python": sys.version.replace("\n", " "),
        "platform": platform.platform(),
        "streamlit": _pkg_version("streamlit"),
        "plotly": _pkg_version("plotly"),
        "pandas": _pkg_version("pandas"),
        "numpy": _pkg_version("numpy"),
        "pytest": _pkg_version("pytest"),
        "python-docx": _pkg_version("python-docx"),
    }
    details = "\n".join([f"{k}: {v}" for k, v in versions.items()])
    return CheckResult("Environment", "PASS", details)


def check_thi_network(project_root: Path) -> CheckResult:
    """Check reachability of the host required for live PV data."""
    host = os.environ.get("THI_PV_HOST", "jupyterhub-wi")

    try:
        import socket

        try:
            ip = socket.gethostbyname(host)
        except Exception as e:
            return CheckResult(
                "THI network/VPN reachability",
                "WARN",
                "Cannot resolve host for live PV data.\n"
                f"Host: {host}\n"
                f"DNS error: {e}\n\n"
                "If you are off-campus, connect via THI VPN and re-run this validation.",
            )

        # Lightweight FTP probe
        try:
            import ftplib

            ftp = ftplib.FTP(host, timeout=8)
            try:
                ftp.login("ftpuser", "ftpuser123")
                ftp.cwd("pvdaten")
                buf = io.BytesIO()
                ftp.retrbinary("RETR pv.csv", buf.write)
                sample = buf.getvalue()[:400].decode("utf-8", errors="replace")
            finally:
                try:
                    ftp.quit()
                except Exception:
                    pass

            preview = "\n".join([l for l in sample.splitlines()[:3] if l.strip()])
            return CheckResult(
                "THI network/VPN reachability",
                "PASS",
                f"Resolved {host} → {ip}\nFTP download OK (pv.csv in memory)\nPreview:\n{preview}",
            )
        except Exception as e:
            return CheckResult(
                "THI network/VPN reachability",
                "WARN",
                f"Resolved {host} → {ip} but FTP access failed.\nError: {e}\n\n"
                "If you are on THI VPN, check credentials and server availability.",
            )

    except Exception as e:
        return CheckResult("THI network/VPN reachability", "WARN", repr(e))


def check_compileall(project_root: Path) -> CheckResult:
    targets = [
        "app.py",
        "dashboard_core.py",
        "ftp_monitor.py",
        "list_ftp.py",
    ]
    existing = [t for t in targets if (project_root / t).exists()]
    cmd = [sys.executable, "-m", "compileall", "-q", *existing]
    code, out = _run(cmd, cwd=project_root, timeout_s=120)
    status = "PASS" if code == 0 else "FAIL"
    return CheckResult("Python bytecode compile", status, out.strip() or "OK")


def check_pytest(project_root: Path) -> CheckResult:
    cmd = [sys.executable, "-m", "pytest", "-q"]
    code, out = _run(cmd, cwd=project_root, timeout_s=180)
    status = "PASS" if code == 0 else "FAIL"
    return CheckResult("Unit tests (pytest)", status, out.strip())


def check_streamlit_smoke(project_root: Path) -> CheckResult:
    """Starts Streamlit briefly to verify the app can boot."""
    app_path = project_root / "app.py"
    if not app_path.exists():
        return CheckResult("Streamlit smoke start", "SKIP", "app.py not found")

    cmd = [
        sys.executable,
        "-m",
        "streamlit",
        "run",
        str(app_path),
        "--server.headless",
        "true",
        "--server.port",
        "8502",
        "--server.address",
        "127.0.0.1",
        "--server.runOnSave",
        "false",
    ]

    try:
        proc = subprocess.Popen(
            cmd,
            cwd=str(project_root),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
    except FileNotFoundError as e:
        return CheckResult("Streamlit smoke start", "FAIL", str(e))

    output_lines: list[str] = []
    try:
        deadline = datetime.now().timestamp() + 12
        while datetime.now().timestamp() < deadline:
            if proc.poll() is not None:
                break
            if proc.stdout is None:
                break
            line = proc.stdout.readline()
            if line:
                output_lines.append(line.rstrip())
                if "Local URL" in line or "You can now view" in line:
                    break

        out = "\n".join(output_lines).strip()
        saw_marker = any(
            ("Local URL" in l or "You can now view" in l) for l in output_lines
        )
        still_running = proc.poll() is None

        status = "PASS" if (saw_marker or still_running) else "FAIL"
        details = out or "(no output captured)"
        return CheckResult("Streamlit smoke start", status, _shorten(details, limit=8000))
    finally:
        try:
            proc.terminate()
            proc.wait(timeout=5)
        except Exception:
            try:
                proc.kill()
            except Exception:
                pass


def check_pv_csv(project_root: Path) -> CheckResult:
    csv_path = project_root / "pv.csv"
    if not csv_path.exists():
        return CheckResult("pv.csv schema", "WARN", "pv.csv not found in repo")

    try:
        import pandas as pd

        df = pd.read_csv(csv_path)
        missing = [c for c in ["timestamp", "total_energy_kWh"] if c not in df.columns]
        ptot = [c for c in df.columns if c.startswith("energy_ptot_") and c.endswith("_kWh")]

        if missing:
            return CheckResult(
                "pv.csv schema",
                "FAIL",
                f"Missing columns: {missing}\nColumns: {list(df.columns)[:30]}...",
            )
        if not ptot:
            return CheckResult(
                "pv.csv schema",
                "FAIL",
                "No energy_ptot_*_kWh columns found",
            )

        return CheckResult(
            "pv.csv schema",
            "PASS",
            f"Rows: {len(df)}\nptot columns: {len(ptot)}\nExample timestamp: {df.iloc[-1]['timestamp']}",
        )
    except Exception as e:
        return CheckResult("pv.csv schema", "FAIL", repr(e))


def _render_html(
    results: Sequence[CheckResult],
    *,
    project_root: Path,
    generated_at: datetime,
) -> str:
    overall = _overall_status(results)
    passed = sum(1 for r in results if r.status.upper() == "PASS")
    warned = sum(1 for r in results if r.status.upper() == "WARN")
    failed = sum(1 for r in results if r.status.upper() == "FAIL")
    skipped = sum(1 for r in results if r.status.upper() == "SKIP")

    _, git_commit = _try_run_git(project_root, ["rev-parse", "--short", "HEAD"])
    ok2, git_status = _try_run_git(project_root, ["status", "--porcelain"])
    dirty = bool(git_status.strip()) if ok2 else False

    def cls(status: str) -> str:
        s = status.upper()
        if s == "PASS":
            return "st st-pass"
        if s == "FAIL":
            return "st st-fail"
        if s == "WARN":
            return "st st-warn"
        if s == "SKIP":
            return "st st-skip"
        return "st"

    rows = []
    for r in sorted(results, key=lambda x: (STATUS_ORDER.get(x.status.upper(), 99), x.name)):
        rows.append(
            "<tr>"
            f"<td>{html.escape(r.name)}</td>"
            f"<td><span class=\"{cls(r.status)}\">{html.escape(r.status.upper())}</span></td>"
            f"<td>{html.escape(_first_line(r.details))}</td>"
            "</tr>"
        )

    detail_blocks = []
    for r in sorted(results, key=lambda x: (STATUS_ORDER.get(x.status.upper(), 99), x.name)):
        detail_blocks.append(
            "<section class=\"card\">"
            f"<h3>{html.escape(r.name)} <span class=\"{cls(r.status)}\">{html.escape(r.status.upper())}</span></h3>"
            f"<pre>{html.escape(_shorten(r.details or '(no details)', limit=16000))}</pre>"
            "</section>"
        )

    dirty_html = " | <strong>Dirty:</strong> yes" if dirty else ""

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>THI Solar Dashboard — Test &amp; Validation Report</title>
  <style>
    :root {{
      --bg: #0b1220;
      --panel: #101a2e;
      --text: #e9eefc;
      --muted: #b6c0dd;
      --border: #233055;
    }}
    body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, Segoe UI, Roboto, Helvetica, Arial, sans-serif; background: var(--bg); color: var(--text); }}
    header {{ padding: 28px 18px 10px; max-width: 1100px; margin: 0 auto; }}
    h1 {{ margin: 0 0 6px; font-size: 22px; }}
    h2 {{ margin: 0 0 10px; font-size: 18px; }}
    h3 {{ margin: 0 0 10px; font-size: 16px; }}
    .meta {{ color: var(--muted); font-size: 13px; line-height: 1.5; }}
    .container {{ max-width: 1100px; margin: 0 auto; padding: 10px 18px 40px; }}
    .grid {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 10px; margin: 14px 0 18px; }}
    .kpi {{ background: var(--panel); border: 1px solid var(--border); border-radius: 12px; padding: 12px 12px; }}
    .kpi .label {{ color: var(--muted); font-size: 12px; }}
    .kpi .value {{ font-size: 18px; margin-top: 4px; }}
    .card {{ background: var(--panel); border: 1px solid var(--border); border-radius: 12px; padding: 14px; margin: 12px 0; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ text-align: left; padding: 10px 8px; border-bottom: 1px solid var(--border); vertical-align: top; }}
    th {{ color: var(--muted); font-weight: 600; font-size: 12px; letter-spacing: .02em; text-transform: uppercase; }}
    pre {{ margin: 10px 0 0; white-space: pre-wrap; word-break: break-word; color: #d8e2ff; background: rgba(0,0,0,.18); padding: 10px; border-radius: 10px; border: 1px solid rgba(255,255,255,.06); }}
    .st {{ display: inline-block; padding: 2px 8px; border-radius: 999px; font-size: 12px; font-weight: 700; border: 1px solid rgba(255,255,255,.12); }}
    .st-pass {{ background: rgba(31,139,76,.18); color: #9ff0c2; border-color: rgba(31,139,76,.35); }}
    .st-warn {{ background: rgba(183,121,31,.18); color: #ffe2b0; border-color: rgba(183,121,31,.35); }}
    .st-fail {{ background: rgba(197,48,48,.18); color: #ffb3b3; border-color: rgba(197,48,48,.35); }}
    .st-skip {{ background: rgba(74,85,104,.18); color: #d6dbe7; border-color: rgba(74,85,104,.35); }}
    .note {{ color: var(--muted); font-size: 13px; }}
    @media (max-width: 860px) {{ .grid {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }} }}
  </style>
</head>
<body>
  <header>
    <h1>THI Solar Dashboard — Test &amp; Validation Report</h1>
    <div class="meta">
      <div><strong>Generated:</strong> {html.escape(generated_at.strftime('%Y-%m-%d %H:%M:%S'))}</div>
      <div><strong>Project:</strong> {html.escape(project_root.name)} | <strong>Git commit:</strong> {html.escape(git_commit)}{dirty_html}</div>
      <div><strong>Overall result:</strong> <span class="{cls(overall)}">{html.escape(overall)}</span></div>
      <div><strong>Python:</strong> {html.escape(sys.executable)}</div>
    </div>
  </header>
  <div class="container">
    <div class="grid">
      <div class="kpi"><div class="label">Checks</div><div class="value">{len(results)}</div></div>
      <div class="kpi"><div class="label">Pass</div><div class="value">{passed}</div></div>
      <div class="kpi"><div class="label">Warn</div><div class="value">{warned}</div></div>
      <div class="kpi"><div class="label">Fail / Skip</div><div class="value">{failed} / {skipped}</div></div>
    </div>

    <section class="card">
      <h2>Summary</h2>
      <table>
        <thead><tr><th>Check</th><th>Status</th><th>Notes</th></tr></thead>
        <tbody>
          {''.join(rows)}
        </tbody>
      </table>
      <p class="note" style="margin:10px 0 0">✨ Share this HTML file for easy viewing (no Word required).</p>
    </section>

    <h2 style="margin-top:18px">Detailed results</h2>
    {''.join(detail_blocks)}

    <section class="card">
      <h2>Manual validation checklist</h2>
      <ul class="note" style="margin:0; padding-left: 18px">
        <li>Start the app: <code>streamlit run app.py</code> and verify it loads.</li>
        <li>If off THI network/VPN: verify the app shows clear guidance/error.</li>
        <li>If on THI network/VPN: verify live values populate and charts render.</li>
        <li>Open the Historical Data view: verify rows load and downloads work.</li>
        <li>Enable auto-refresh for 1–2 minutes: verify no obvious CPU/memory spikes.</li>
      </ul>
    </section>
  </div>
</body>
</html>"""


def _render_markdown(
    results: Sequence[CheckResult],
    *,
    project_root: Path,
    generated_at: datetime,
) -> str:
    overall = _overall_status(results)
    passed = sum(1 for r in results if r.status.upper() == "PASS")
    warned = sum(1 for r in results if r.status.upper() == "WARN")
    failed = sum(1 for r in results if r.status.upper() == "FAIL")
    skipped = sum(1 for r in results if r.status.upper() == "SKIP")

    ok1, git_commit = _try_run_git(project_root, ["rev-parse", "--short", "HEAD"])
    ok2, git_status = _try_run_git(project_root, ["status", "--porcelain"])
    dirty = bool(git_status.strip()) if ok2 else False
    if not ok1:
        git_commit = "(not available)"

    lines: list[str] = []
    lines.append("# THI Solar Dashboard — Test & Validation Report")
    lines.append("")
    lines.append(f"**Generated:** {generated_at.strftime('%Y-%m-%d %H:%M:%S')}  ")
    lines.append(f"**Project:** {project_root.name}  ")
    lines.append(f"**Overall result:** {_status_badge(overall)}  ")
    lines.append("")
    lines.append("## Executive summary")
    lines.append("")
    lines.append(
        f"- Checks: {len(results)} total — {passed} pass, {warned} warn, {failed} fail, {skipped} skip"
    )
    lines.append(f"- Git commit: {git_commit}{' (dirty working tree)' if dirty else ''}")
    lines.append(f"- Python: `{sys.executable}`")

    lines.append("")
    lines.append("## Summary table")
    lines.append("")
    lines.append("| Check | Status | Notes |")
    lines.append("|---|---:|---|")
    for r in sorted(results, key=lambda x: (STATUS_ORDER.get(x.status.upper(), 99), x.name)):
        lines.append(f"| {r.name} | {_status_badge(r.status)} | {_first_line(r.details)} |")

    lines.append("")
    lines.append("## Detailed results")
    for r in sorted(results, key=lambda x: (STATUS_ORDER.get(x.status.upper(), 99), x.name)):
        lines.append("")
        lines.append(f"### {r.name} — {_status_badge(r.status)}")
        lines.append("")
        if r.details:
            lines.append("```text")
            lines.append(_shorten(r.details, limit=12000).rstrip())
            lines.append("```")
        else:
            lines.append("(no details)")

    lines.append("")
    return "\n".join(lines)


def write_docx(
    results: Sequence[CheckResult],
    *,
    project_root: Path,
    out_path: Path,
) -> None:
    doc = Document()

    doc.add_heading("THI Solar Dashboard — Test & Validation Report", level=0)

    generated_at = datetime.now()
    overall = _overall_status(results)
    doc.add_paragraph(f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Project: {project_root.name}")
    doc.add_paragraph(f"Overall result: {overall}")
    doc.add_paragraph(f"Python executable: {sys.executable}")

    ok1, git_commit = _try_run_git(project_root, ["rev-parse", "--short", "HEAD"])
    ok2, git_status = _try_run_git(project_root, ["status", "--porcelain"])
    if ok1:
        dirty = " (dirty working tree)" if (ok2 and git_status.strip()) else ""
        doc.add_paragraph(f"Git commit: {git_commit}{dirty}")

    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    hdr = table.rows[0].cells
    hdr[0].text = "Check"
    hdr[1].text = "Status"
    hdr[2].text = "Notes"

    for r in results:
        row = table.add_row().cells
        row[0].text = r.name
        row[1].text = r.status
        row[2].text = _first_line(r.details, limit=140)

    doc.add_heading("Detailed Results", level=1)

    for r in results:
        doc.add_heading(f"{r.name} — {r.status}", level=2)
        if r.details:
            doc.add_paragraph(r.details)

    doc.add_paragraph("")
    doc.add_paragraph("End of report.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="THI Solar Dashboard — Test & Validation (single script)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python validate.py                # Generate all formats
  python validate.py --html-only    # HTML only (recommended)
  python validate.py --quick        # Skip Streamlit smoke test
        """,
    )
    parser.add_argument(
        "--html-only",
        action="store_true",
        help="Generate HTML report only (fastest, easiest to share)",
    )
    parser.add_argument(
        "--quick",
        action="store_true",
        help="Skip slow checks (Streamlit smoke test)",
    )
    parser.add_argument(
        "--skip-network",
        action="store_true",
        help="Skip THI network/VPN reachability check",
    )
    args = parser.parse_args()

    project_root = Path(__file__).resolve().parent
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = project_root / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("THI Solar Dashboard — Test & Validation")
    print("=" * 60)
    print()

    results: list[CheckResult] = []

    print("⏳ Checking environment...")
    results.append(check_environment())

    if not args.skip_network:
        print("⏳ Checking THI network/VPN reachability...")
        results.append(check_thi_network(project_root))

    print("⏳ Compiling Python files...")
    results.append(check_compileall(project_root))

    print("⏳ Checking pv.csv schema...")
    results.append(check_pv_csv(project_root))

    print("⏳ Running unit tests (pytest)...")
    results.append(check_pytest(project_root))

    if args.quick:
        print("⏭️  Skipping Streamlit smoke test (--quick)")
        results.append(CheckResult("Streamlit smoke start", "SKIP", "Skipped by --quick"))
    else:
        print("⏳ Starting Streamlit smoke test (this may take ~15 seconds)...")
        results.append(check_streamlit_smoke(project_root))

    print()
    print("=" * 60)
    print("Generating report...")
    print("=" * 60)
    print()

    generated_at = datetime.now()
    wrote: list[Path] = []

    if args.html_only:
        out_path = out_dir / f"validation_report_{timestamp}.html"
        out_path.write_text(
            _render_html(results, project_root=project_root, generated_at=generated_at),
            encoding="utf-8",
        )
        wrote.append(out_path)
    else:
        # Generate all formats
        out_path = out_dir / f"validation_report_{timestamp}.html"
        out_path.write_text(
            _render_html(results, project_root=project_root, generated_at=generated_at),
            encoding="utf-8",
        )
        wrote.append(out_path)

        out_path = out_dir / f"validation_report_{timestamp}.md"
        out_path.write_text(
            _render_markdown(results, project_root=project_root, generated_at=generated_at),
            encoding="utf-8",
        )
        wrote.append(out_path)

        out_path = out_dir / f"validation_report_{timestamp}.docx"
        write_docx(results, project_root=project_root, out_path=out_path)
        wrote.append(out_path)

    print("✅ Report files created:")
    for p in wrote:
        print(f"   📄 {p.relative_to(project_root)}")
    print()

    # Summary
    overall = _overall_status(results)
    passed = sum(1 for r in results if r.status.upper() == "PASS")
    warned = sum(1 for r in results if r.status.upper() == "WARN")
    failed = sum(1 for r in results if r.status.upper() == "FAIL")

    print("=" * 60)
    print(f"Overall result: {_status_badge(overall)}")
    print(f"Checks: {len(results)} total — {passed} pass, {warned} warn, {failed} fail")
    print("=" * 60)
    print()

    if args.html_only:
        print("💡 Tip: Open the HTML file in your browser to view the report.")
    else:
        print("💡 Tip: Share the HTML file with your professor (no Word needed).")

    # Exit non-zero if any FAIL
    failed_checks = [r for r in results if r.status == "FAIL"]
    return 1 if failed_checks else 0


if __name__ == "__main__":
    raise SystemExit(main())
